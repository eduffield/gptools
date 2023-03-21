import openai
from canvasapi import Canvas
from bs4 import BeautifulSoup
import html5lib
from docx import Document
from docx.shared import Pt

# Proof of concept for a "Homework Machine" using
# Canvas' REST api and gpt-3.5. For BACS 385 use only.

openai.api_key = "PASTE_OPENAI_TOKEN"
canvas_key = "PASTE_CANVAS_TOKEN"
canvas_url = "https://unco.instructure.com"

def listmodels():
    models = openai.Model.list()
    model_names = [model.id for model in models["data"]]
    for model in model_names:
        print(model)

def query(userprompt):
    response = openai.Completion.create(
                engine="text-davinci-003",
                prompt=userprompt,
                max_tokens=1024,
                n=1,
                stop=None,
                temperature=0.5)
    return response.choices[0].text.strip()

def getassignment():
    canvas = Canvas(canvas_url, canvas_key)
    course = canvas.get_course(95764)
    print(course.name)
    assigns = course.get_assignments()

    largest = 0
    for assign in assigns:
        num = assign.id
        if largest < num:
            largest = num
    print("Latest Assignment: " + str(course.get_assignment(largest).name))
    html = course.get_assignment(largest).description
    s = BeautifulSoup(html, "html5lib")
    fin = s.find("p").text
    print(fin)
    return (course.name, str(course.get_assignment(largest).name), fin)



def writepaper(classname, activity, promptforbot):
    names = ["Dillon Blair, Bennett Wiley,",
             "Anson Cordeiro, Evan Duffield,",
             "Ezikiel Quinones"]

    d = Document()
    sty = d.styles['Normal']
    font = sty.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    sty.paragraph_format.line_spacing = 1.5



    d.add_paragraph(classname)
    d.add_paragraph(activity)
    [d.add_paragraph(x) for x in names]
    d.add_paragraph(" ")
    d.add_paragraph(" ")
    d.add_paragraph(query(promptforbot))

    d.save(activity + ".docx")

if __name__ == "__main__":
    assign = getassignment()
    writepaper(assign[0], assign[1], assign[2])